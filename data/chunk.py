"""
chunk_laws.py — Chia chunk văn bản pháp luật từ file .doc/.docx cho RAG

Xử lý trực tiếp .doc bằng LibreOffice (convert sang .txt trong bộ nhớ tạm).
Không cần chuyển file sang .docx.

Cấu trúc thư mục input:
    data/law_clean/
        amended/          # Luật đã được sửa đổi
        in-force/         # Luật đang còn hiệu lực
        partially-expire/ # Luật hết hiệu lực 1 phần

Output:
    data/rag_chunks/
        rag_corpus.jsonl  — mỗi dòng là 1 chunk, sẵn để embed
        _stats.json       — thống kê tổng quan

Cài đặt:
    pip install tqdm
    (LibreOffice cần có sẵn: sudo apt install libreoffice)
"""

import hashlib
import json
import os
import re
import subprocess
import tempfile
import time
from pathlib import Path

# ─── CẤU HÌNH ─────────────────────────────────────────────────────────────────
DATA_DIR   = Path("law-clean")
OUTPUT_DIR = Path("data/rag_chunks")
MAX_CHUNK_CHARS = 1500   # Độ dài tối đa mỗi chunk (ký tự)
OVERLAP_CHARS   = 150    # Overlap giữa các chunk

# Map tên thư mục → status metadata
STATUS_MAP = {
    "amended":          "da_sua_doi",
    "in-force":         "dang_hieu_luc",
    "partially-expire": "het_hieu_luc_mot_phan",
}
# ──────────────────────────────────────────────────────────────────────────────

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# ══════════════════════════════════════════════════════════════════════════════
# BƯỚC 1: ĐỌC FILE .doc / .docx → plain text
# ══════════════════════════════════════════════════════════════════════════════

def doc_to_text(filepath: Path) -> str:
    """
    Đọc .doc hoặc .docx → plain text.
    Pipeline: .doc → .docx (LibreOffice) → text (python-docx)
    Tránh hoàn toàn vấn đề encoding của LibreOffice trên Windows.
    """
    import docx

    suffix = filepath.suffix.lower()

    if suffix == ".docx":
        docx_path = filepath
        tmp_dir = None
    else:
        # Convert .doc → .docx vào thư mục tạm
        tmp_dir = tempfile.mkdtemp()
        soffice = _find_soffice()
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "docx",
             str(filepath), "--outdir", tmp_dir],
            capture_output=True, timeout=60
        )
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice lỗi: {result.stderr}")

        docx_files = list(Path(tmp_dir).glob("*.docx"))
        if not docx_files:
            raise RuntimeError(f"Không tìm thấy file .docx sau convert: {filepath}")
        docx_path = docx_files[0]

    # Đọc bằng python-docx (luôn UTF-8, không phụ thuộc hệ thống)
    doc = docx.Document(str(docx_path))
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)
    # Đọc thêm text trong bảng (header luật thường nằm trong table)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        lines.append(para.text)

    # Dọn thư mục tạm
    if tmp_dir:
        import shutil
        shutil.rmtree(tmp_dir, ignore_errors=True)

    return "\n".join(lines)


def _find_soffice() -> str:
    """Tìm đường dẫn LibreOffice trên Windows / Linux / Mac."""
    import shutil

    # Thử PATH trước (Linux/Mac)
    found = shutil.which("soffice")
    if found:
        return found

    # Các đường dẫn phổ biến trên Windows
    windows_paths = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for p in windows_paths:
        if Path(p).exists():
            return p

    raise RuntimeError(
        "Không tìm thấy LibreOffice.\n"
        "Hãy cài tại: https://www.libreoffice.org/download/\n"
        "Hoặc thêm thư mục 'program' của LibreOffice vào PATH."
    )
# ══════════════════════════════════════════════════════════════════════════════
# BƯỚC 2: EXTRACT METADATA TỪ PHẦN ĐẦU FILE
# ══════════════════════════════════════════════════════════════════════════════

def extract_metadata(text: str, filepath: Path, status: str) -> dict:
    """
    Trích xuất metadata từ phần header văn bản pháp luật VN.

    Ví dụ header:
        QUỐC HỘI
        Luật số: 01/2021/QH15
        LUẬT
        SỬA ĐỔI, BỔ SUNG MỘT SỐ ĐIỀU...
        Căn cứ Hiến pháp...
        Quốc hội ban hành...
    """
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    meta = {
        "filename":        filepath.name,
        "status":          status,
        "so_hieu":         None,   # VD: 01/2021/QH15
        "ten_luat":        None,   # Tên đầy đủ
        "co_quan_bh":      None,   # Cơ quan ban hành
        "ngay_thong_qua":  None,   # Ngày thông qua
        "ngay_hieu_luc":   None,   # Ngày có hiệu lực
        "nguoi_ky":        None,   # Người ký
        "loai_van_ban":    None,   # LUẬT / NGHỊ QUYẾT / v.v.
        "sua_doi_luat":    None,   # Luật gốc bị sửa đổi (nếu có)
    }

    # ── Số hiệu ──────────────────────────────────────────────────────────────
    # Pattern: "Luật số: 01/2021/QH15" hoặc "Số: 01/2021/QH15"
   # Bắt cả: Luật-02-2021-QH15, Bộ-luật-91-2015-QH13
    m_file = re.search(
        r"(\d+)-(\d{4})-(QH\d+|UBTVQH\d+|CP|TTg)",
        filepath.stem, re.IGNORECASE
    )
    if m_file:
        meta["so_hieu"] = f"{m_file.group(1)}/{m_file.group(2)}/{m_file.group(3).upper()}"
    else:
        # Fallback: tìm trong 10 dòng đầu, chỉ match dòng ngắn (không phải câu văn)
        for line in lines[:10]:
            m = re.match(r"^(?:Luật số|Số)\s*[:\-]\s*(\d+/\d{4}/[\w\-]+)\s*$", line.strip(), re.IGNORECASE)
            if m:
                meta["so_hieu"] = m.group(1).strip()
                break
      # ── Tên luật — dòng IN HOA dài ngay sau "LUẬT" / "BỘ LUẬT" ──────────
    loai_kw = ["LUẬT", "BỘ LUẬT", "NGHỊ QUYẾT", "NGHỊ ĐỊNH", "THÔNG TƯ", "PHÁP LỆNH"]
    found_type = False
    for line in lines[:40]:
        stripped = line.strip()
        if found_type:
            # Bỏ qua dòng trống hoặc dòng trang trí
            if not stripped or re.match(r"^[_\-\s]+$", stripped):
                continue
            # Tên luật thường IN HOA hoặc dài > 10 ký tự
            if len(stripped) > 10:
                meta["ten_luat"] = stripped
                break
        if stripped.upper() in loai_kw or re.match(r"^(LUẬT|BỘ LUẬT)\s*$", stripped, re.IGNORECASE):
            found_type = True
    # ── Cơ quan ban hành ─────────────────────────────────────────────────────
    co_quan_kw = ["QUỐC HỘI", "CHÍNH PHỦ", "THỦ TƯỚNG", "BỘ ", "ỦY BAN"]
    for line in lines[:10]:
        if any(line.upper().startswith(kw) for kw in co_quan_kw):
            meta["co_quan_bh"] = line
            break

    # ── Loại văn bản ─────────────────────────────────────────────────────────
    loai_kw = ["LUẬT", "NGHỊ QUYẾT", "NGHỊ ĐỊNH", "THÔNG TƯ", "QUYẾT ĐỊNH", "PHÁP LỆNH"]
    for line in lines[:20]:
        if line.upper() in loai_kw:
            meta["loai_van_ban"] = line.upper()
            break

    # ── Tên luật: dòng CAPS dài sau loại văn bản ─────────────────────────────
    found_type = False
    for line in lines[:30]:
        # Chỉ match dòng CÓ "Luật số:" hoặc "Số:" đứng riêng, không phải giữa câu
        m = re.match(r"^(?:Luật số|Số)\s*[:\-]\s*(\d+/\d{4}/[\w\-]+)", line.strip(), re.IGNORECASE)
        if m:
            meta["so_hieu"] = m.group(1).strip()
            break
        if line.upper() in loai_kw:
            found_type = True

    # ── Luật gốc bị sửa đổi ──────────────────────────────────────────────────
    # Pattern: "của Luật X số YY/20ZZ/QHWW"
    for line in lines[:30]:
        m = re.search(
            r"của\s+(Luật\s+[\w\s]+?)\s+số\s+([\d/A-Z]+)",
            line, re.IGNORECASE
        )
        if m:
            meta["sua_doi_luat"] = f"{m.group(1).strip()} số {m.group(2).strip()}"
            break

    # ── Ngày thông qua ────────────────────────────────────────────────────────
    # Pattern: "thông qua ngày DD tháng MM năm YYYY"
    for line in lines:
        m = re.search(
            r"thông qua ngày\s+(\d{1,2})\s+tháng\s+(\d{1,2})\s+năm\s+(\d{4})",
            line, re.IGNORECASE
        )
        if m:
            meta["ngay_thong_qua"] = f"{m.group(1)}/{m.group(2)}/{m.group(3)}"
            break

    # ── Ngày hiệu lực ─────────────────────────────────────────────────────────
    # Pattern: "có hiệu lực thi hành từ ngày DD tháng MM năm YYYY"
    for line in lines:
        m = re.search(
            r"có hiệu lực.*?ngày\s+(\d{1,2})\s+tháng\s+(\d{1,2})\s+năm\s+(\d{4})",
            line, re.IGNORECASE
        )
        if m:
            meta["ngay_hieu_luc"] = f"{m.group(1)}/{m.group(2)}/{m.group(3)}"
            break

    # ── Người ký ──────────────────────────────────────────────────────────────
    for i, line in enumerate(lines):
        if "ĐÃ KÝ" in line.upper() or "Đã ký" in line:
            # Lấy dòng tiếp theo sau "Đã ký:"
            m = re.search(r"(?:Đã ký|ĐÃ KÝ)\s*[:\-]?\s*(.+)", line, re.IGNORECASE)
            if m and m.group(1).strip():
                meta["nguoi_ky"] = m.group(1).strip()
            elif i + 1 < len(lines):
                meta["nguoi_ky"] = lines[i + 1]
            break

    return meta


# ══════════════════════════════════════════════════════════════════════════════
# BƯỚC 3: TÁCH CÁC PHẦN CẤU TRÚC (Điều, Chương, Phụ lục)
# ══════════════════════════════════════════════════════════════════════════════

# Regex nhận diện các boundary quan trọng
RE_DIEU    = re.compile(r"^(Điều\s+\d+[\.\:]?\s*.{0,120})$", re.IGNORECASE)
RE_CHUONG  = re.compile(r"^(Chương\s+[IVXLCDM\d]+[\.\:]?\s*.{0,120})$", re.IGNORECASE)
RE_MUC     = re.compile(r"^(Mục\s+\d+[\.\:]?\s*.{0,120})$", re.IGNORECASE)
RE_PHUCLUC = re.compile(r"^(PHỤ LỤC.{0,120})$", re.IGNORECASE)
RE_HEADER  = re.compile(
    r"^(?:QUỐC HỘI|CHÍNH PHỦ|CỘNG HÒA XÃ HỘI|Độc lập|________|\d{1,3}$|Số thứ tự|Mã số)",
    re.IGNORECASE
)


def parse_sections(text: str) -> list[dict]:
    """
    Tách văn bản thành các section có cấu trúc:
      - Mở đầu (căn cứ, lời nói đầu)
      - Từng Điều luật
      - Phụ lục (nếu có)

    Trả về list[{"section_type", "heading", "content"}]
    """
    lines = text.splitlines()
    sections = []
    current = {"section_type": "mo_dau", "heading": "Mở đầu", "content": []}
    current_chuong = ""
    current_muc = ""

    # Bỏ qua phần header (trước "Căn cứ" hoặc trước Điều 1)
    start_idx = 0
    for i, line in enumerate(lines):
        stripped = line.strip()
        if RE_HEADER.match(stripped) or len(stripped) < 3:
            continue
        if re.match(r"^Căn cứ", stripped, re.IGNORECASE) or RE_DIEU.match(stripped):
            start_idx = i
            break

    for line in lines[start_idx:]:
        stripped = line.strip()
        if not stripped:
            current["content"].append("")
            continue

        # ── Chương mới ──
        if RE_CHUONG.match(stripped):
            current_chuong = stripped
            current_muc = ""
            current["content"].append(stripped)
            continue

        # ── Mục mới ──
        if RE_MUC.match(stripped):
            current_muc = stripped
            current["content"].append(stripped)
            continue

        # ── Điều mới → lưu section hiện tại, mở section mới ──
        if RE_DIEU.match(stripped):
            if current["content"]:
                sections.append({**current,
                    "content": "\n".join(current["content"]).strip()})
            current = {
                "section_type": "dieu",
                "heading":      stripped,
                "chuong":       current_chuong,
                "muc":          current_muc,
                "content":      [],
            }
            continue

        # ── Phụ lục ──
        if RE_PHUCLUC.match(stripped):
            if current["content"]:
                sections.append({**current,
                    "content": "\n".join(current["content"]).strip()})
            current = {
                "section_type": "phu_luc",
                "heading":      stripped,
                "chuong":       "",
                "muc":          "",
                "content":      [],
            }
            continue

        current["content"].append(stripped)

    # Lưu section cuối
    if current["content"]:
        sections.append({**current,
            "content": "\n".join(current["content"]).strip()})

    # Lọc bỏ section rỗng hoặc chỉ là chữ ký / kết thúc
    result = []
    for s in sections:
        if len(s["content"]) < 20:
            continue
        result.append(s)

    return result


# ══════════════════════════════════════════════════════════════════════════════
# BƯỚC 4: CHUNK — chia section dài thành các chunk nhỏ hơn
# ══════════════════════════════════════════════════════════════════════════════

def split_into_chunks(text: str,
                      max_chars: int = MAX_CHUNK_CHARS,
                      overlap: int = OVERLAP_CHARS) -> list[str]:
    """Chia text dài thành chunks, cắt tại ranh giới câu/đoạn."""
    if len(text) <= max_chars:
        return [text]

    chunks = []
    start = 0
    while start < len(text):
        end = start + max_chars
        if end >= len(text):
            chunks.append(text[start:].strip())
            break

        # Ưu tiên cắt tại: xuống dòng > ". " > " "
        cut = -1
        for sep in ("\n", ". ", " "):
            pos = text.rfind(sep, start + overlap, end)
            if pos > start:
                cut = pos + len(sep)
                break
        if cut == -1:
            cut = end

        chunks.append(text[start:cut].strip())
        start = cut - overlap   # overlap để không mất ngữ cảnh

    return [c for c in chunks if c]


# ══════════════════════════════════════════════════════════════════════════════
# BƯỚC 5: TẠO CHUNK RECORDS đầy đủ metadata
# ══════════════════════════════════════════════════════════════════════════════

def make_chunks(sections: list[dict], meta: dict) -> list[dict]:
    """Kết hợp section + metadata → list chunk records cho RAG."""
    chunks = []
    for sec in sections:
        texts = split_into_chunks(sec["content"])
        for i, text in enumerate(texts):
            law_context = f"[{meta['loai_van_ban'] or 'LUẬT'} {meta['so_hieu'] or ''}] {meta['ten_luat'] or ''}\n"
            if sec['section_type'] == 'dieu':
                law_context += f"{sec.get('chuong','')}\n{sec.get('muc','')}\n{sec['heading']}\n".strip() + "\n"
            heading = sec.get("heading", "").strip()
            chunk = {
                # ── Định danh ──────────────────────────────────────────────
                
                "chunk_id": (
                    f"{meta['so_hieu']}"
                    f"__{sec['section_type']}"
                    f"__{hashlib.md5(heading.encode()).hexdigest()[:8]}"
                    f"__{i}"
                ),

                # ── Nội dung ───────────────────────────────────────────────
                # Tạo prefix context cho mỗi chunk
                

                # Ghép vào text
                "text": law_context + text,
                "char_len": len(text),

                # ── Vị trí trong văn bản ───────────────────────────────────
                "section_type": sec["section_type"],   # dieu / mo_dau / phu_luc
                "heading":      sec["heading"],
                "chuong":       sec.get("chuong", ""),
                "muc":          sec.get("muc", ""),
                "chunk_idx":    i,
                "total_chunks": len(texts),

                # ── Metadata văn bản ───────────────────────────────────────
                "so_hieu":        meta["so_hieu"],
                "ten_luat":       meta["ten_luat"],
                "loai_van_ban":   meta["loai_van_ban"],
                "co_quan_bh":     meta["co_quan_bh"],
                "ngay_thong_qua": meta["ngay_thong_qua"],
                "ngay_hieu_luc":  meta["ngay_hieu_luc"],
                "nguoi_ky":       meta["nguoi_ky"],
                "sua_doi_luat":   meta["sua_doi_luat"],

                # ── Trạng thái hiệu lực (từ thư mục) ─────────────────────
                "status":   meta["status"],   # dang_hieu_luc | da_sua_doi | het_hieu_luc_mot_phan
                "filename": meta["filename"],
            }
            chunks.append(chunk)
    # Gộp chunk quá ngắn vào chunk trước
    merged = []
    for chunk in chunks:
        if merged and chunk["char_len"] < 100 and chunk["section_type"] == merged[-1]["section_type"]:
            merged[-1]["text"]     += "\n" + chunk["text"]
            merged[-1]["char_len"] += chunk["char_len"]
            merged[-1]["total_chunks"] -= 1
        else:
            merged.append(chunk)
    return merged


# ══════════════════════════════════════════════════════════════════════════════
# PIPELINE CHÍNH
# ══════════════════════════════════════════════════════════════════════════════

def process_file(filepath: Path, status: str) -> list[dict] | None:
    """Xử lý 1 file .doc/.docx → list chunk."""
    try:
        text = doc_to_text(filepath)
    except Exception as e:
        print(f"    ✗ Lỗi đọc file: {e}")
        return None

    meta     = extract_metadata(text, filepath, status)
    sections = parse_sections(text)
    chunks   = make_chunks(sections, meta)

    return chunks


def main():
    if not DATA_DIR.exists():
        print(f"✗ Không tìm thấy thư mục: {DATA_DIR}")
        print(f"  Hãy đặt script này ngang hàng với thư mục 'data/'")
        return

    output_file = OUTPUT_DIR / "rag_corpus.jsonl"
    stats_file  = OUTPUT_DIR / "_stats.json"

    total_files  = 0
    total_chunks = 0
    errors       = []
    stats_per_status = {}

    print(f"\n{'='*60}")
    print(f"  Chunk Laws → RAG Corpus")
    print(f"  Input : {DATA_DIR}")
    print(f"  Output: {output_file}")
    print(f"{'='*60}\n")

    with output_file.open("w", encoding="utf-8") as out:
        for folder_name, status in STATUS_MAP.items():
            folder = DATA_DIR / folder_name
            if not folder.exists():
                print(f"  ⚠ Thư mục không tồn tại: {folder} — bỏ qua")
                continue

            doc_files = sorted(
                list(folder.glob("*.doc")) + list(folder.glob("*.docx"))
            )
            print(f"[{folder_name}] {len(doc_files)} file")
            stats_per_status[status] = {"files": len(doc_files), "chunks": 0, "errors": 0}

            for fp in doc_files:
                print(f"  → {fp.name[:60]}")
                chunks = process_file(fp, status)

                if chunks is None:
                    errors.append(str(fp))
                    stats_per_status[status]["errors"] += 1
                    continue

                for chunk in chunks:
                    out.write(json.dumps(chunk, ensure_ascii=False) + "\n")

                total_files  += 1
                total_chunks += len(chunks)
                stats_per_status[status]["chunks"] += len(chunks)
                print(f"     ✓ {len(chunks)} chunks")

    # Ghi stats
    stats = {
        "generated_at":    time.strftime("%Y-%m-%dT%H:%M:%S"),
        "total_files":     total_files,
        "total_chunks":    total_chunks,
        "avg_chunks_per_file": round(total_chunks / total_files, 1) if total_files else 0,
        "max_chunk_chars": MAX_CHUNK_CHARS,
        "overlap_chars":   OVERLAP_CHARS,
        "by_status":       stats_per_status,
        "errors":          errors,
    }
    stats_file.write_text(json.dumps(stats, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"\n{'='*60}")
    print(f"  ✅ Xong: {total_chunks} chunks từ {total_files} file")
    print(f"  Output: {output_file}")
    if errors:
        print(f"  ⚠ {len(errors)} file lỗi — xem {stats_file}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    main()